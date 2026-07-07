from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document

ROOT = Path(r"C:\Users\Administrator\Documents\exam")
APP_FILE = ROOT / "index.html"
REPORT_FILE = ROOT / "培训题导入报告.json"

SOURCES = [
    {
        "path": Path(r"C:\Users\Administrator\Desktop\知识点\01课件整理\梁宁产品思维30讲_知识点与题库.docx"),
        "business": "product-thinking",
    },
    {
        "path": Path(r"C:\Users\Administrator\Desktop\知识点\01课件整理\数据科学AI培训材料V1.0_知识点与题库.docx"),
        "business": "digital-transformation",
    },
    {
        "path": Path(r"C:\Users\Administrator\Desktop\知识点\01课件整理\云原生实操培训课件_知识点与题库.docx"),
        "business": "cloud-native",
    },
]

TYPE_MAP = {
    "单选": "single",
    "单选题": "single",
    "多选": "multiple",
    "多选题": "multiple",
}

QUESTION_START_RE = re.compile(r"^(?:第?\s*\d+\s*[.、]\s*)?(单选题|多选题|单选|多选)\s*(\d+)?\s*[.、．]?\s*(.*)$")
OPTION_RE = re.compile(r"^([A-H])\s*[.、．]\s*(.+)$")
ANSWER_RE = re.compile(r"^(?:答案|正确答案)\s*[:：]\s*(.+)$")
ANALYSIS_RE = re.compile(r"^(?:解析|答案解析)\s*[:：]\s*(.*)$")
SECTION_RE = re.compile(r"^[一二三四五六七八九十]+、")


def clean(text: object) -> str:
    return "" if text is None else str(text).strip()


def normalize_key(text: str) -> str:
    text = re.sub(r"^\s*(?:第?\s*\d+\s*[.、．]\s*)?", "", text)
    text = re.sub(r"^(?:单选题|多选题|单选|多选)\s*\d*\s*[.、．]?\s*", "", text)
    text = re.sub(r"【[^】]+】", "", text)
    return re.sub(r"[\W_]+", "", text).lower()


def find_questions_bounds(html: str) -> tuple[int, int]:
    marker = "    const questions = "
    start = html.find(marker)
    if start < 0:
        raise RuntimeError("未找到 const questions")
    array_start = html.find("[", start)
    depth = 0
    in_string = False
    escape = False
    for idx in range(array_start, len(html)):
        ch = html[idx]
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
        else:
            if ch == '"':
                in_string = True
            elif ch == "[":
                depth += 1
            elif ch == "]":
                depth -= 1
                if depth == 0:
                    semi = html.find(";", idx)
                    return start, semi + 1
    raise RuntimeError("questions 数组未闭合")


def load_questions(html: str) -> list[dict]:
    start, end = find_questions_bounds(html)
    raw = html[start:end].split("=", 1)[1].rsplit(";", 1)[0].strip()
    return json.loads(raw)


def dump_questions(html: str, questions: list[dict]) -> str:
    start, end = find_questions_bounds(html)
    js = "    const questions = " + json.dumps(questions, ensure_ascii=False, indent=6) + ";"
    js = js.replace("\n", "\n    ", 1)
    return html[:start] + js + html[end:]


def docx_lines(path: Path) -> list[str]:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        text = clean(p.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            text = " ".join(cell for cell in cells if cell)
            if text:
                lines.append(text)
    return lines


def split_blocks(lines: list[str]) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    in_question_area = False
    for line in lines:
        if re.match(r"^[一二三四五六七八九十]+、.*(?:单选题|多选题)", line):
            in_question_area = True
            continue
        if SECTION_RE.match(line) and "题" not in line and current:
            blocks.append(current)
            current = []
            in_question_area = False
        start = QUESTION_START_RE.match(line)
        if start:
            if current:
                blocks.append(current)
            current = [line]
            in_question_area = True
        elif current and in_question_area:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def parse_answer(raw: str, qtype: str) -> list[str]:
    raw = clean(raw).upper()
    raw = raw.replace("，", "").replace(",", "").replace("、", "").replace(" ", "")
    letters = re.findall(r"[A-H]", raw)
    if qtype == "single" and len(letters) > 1:
        return letters[:1]
    return letters


def parse_block(block: list[str], source_name: str, business: str) -> dict | None:
    head = block[0]
    match = QUESTION_START_RE.match(head)
    if not match:
        return None
    qtype = TYPE_MAP.get(match.group(1))
    question = clean(match.group(3))
    options: dict[str, str] = {}
    answer = ""
    analysis_parts: list[str] = []
    current_option: str | None = None
    mode = "question"

    for line in block[1:]:
        option_match = OPTION_RE.match(line)
        answer_match = ANSWER_RE.match(line)
        analysis_match = ANALYSIS_RE.match(line)
        next_question = QUESTION_START_RE.match(line)
        if next_question:
            break
        if option_match:
            current_option = option_match.group(1)
            options[current_option] = clean(option_match.group(2))
            mode = "options"
        elif answer_match:
            answer = clean(answer_match.group(1))
            current_option = None
            mode = "answer"
        elif analysis_match:
            text = clean(analysis_match.group(1))
            if text:
                analysis_parts.append(text)
            current_option = None
            mode = "analysis"
        elif mode == "analysis":
            analysis_parts.append(line)
        elif current_option:
            options[current_option] = f"{options[current_option]} {line}".strip()
        elif mode == "question":
            question = f"{question} {line}".strip()

    option_values = [options[key] for key in sorted(options)]
    parsed_answer = parse_answer(answer, qtype or "")
    if qtype not in {"single", "multiple"} or not question or not option_values or not parsed_answer:
        return None
    if any(letter not in options for letter in parsed_answer):
        return None
    return {
        "id": "",
        "type": qtype,
        "business": business,
        "question": question,
        "options": option_values,
        "answer": parsed_answer,
        "analysis": " ".join(analysis_parts).strip() or f"本题答案为{''.join(parsed_answer)}。",
        "source": source_name,
    }


def renumber(questions: list[dict]) -> list[dict]:
    for idx, question in enumerate(questions, 1):
        question["id"] = f"q{idx:04d}"
    return questions


def main() -> None:
    html = APP_FILE.read_text(encoding="utf-8-sig")
    existing = load_questions(html)
    existing_keys = {normalize_key(q.get("question", "")) for q in existing}
    seen_new: set[str] = set()
    imported: list[dict] = []
    skipped: list[dict] = []
    parsed_counts = Counter()
    type_counts = Counter()
    business_counts = Counter()

    for source in SOURCES:
        path = source["path"]
        business = source["business"]
        blocks = split_blocks(docx_lines(path))
        parsed_counts[f"blocks:{path.name}"] = len(blocks)
        for block in blocks:
            item = parse_block(block, path.name, business)
            if not item:
                continue
            parsed_counts[f"parsed:{path.name}"] += 1
            key = normalize_key(item["question"])
            if key in existing_keys or key in seen_new:
                skipped.append({"question": item["question"], "business": item["business"], "source": item["source"]})
                continue
            seen_new.add(key)
            imported.append({key: value for key, value in item.items() if key != "source"})
            type_counts[item["type"]] += 1
            business_counts[item["business"]] += 1

    updated = renumber(existing + imported)
    APP_FILE.write_text(dump_questions(html, updated), encoding="utf-8")
    report = {
        "existing_count": len(existing),
        "candidate_count": sum(value for key, value in parsed_counts.items() if key.startswith("parsed:")),
        "imported_count": len(imported),
        "skipped_duplicate_count": len(skipped),
        "updated_count": len(updated),
        "parsed_counts": dict(parsed_counts),
        "imported_by_type": dict(type_counts),
        "imported_by_business": dict(business_counts),
        "skipped_duplicates_sample": skipped[:80],
    }
    REPORT_FILE.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
